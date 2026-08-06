from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document


SOURCES = [
    Path(
        r"C:\Users\marco\OneDrive - Caja Costarricense de Seguro Social\2025"
        r"\08-Agosto-2025\Aires Acondicionados"
        r"\GF-EIM-IT-0031-2025---Esp. Tec. Aires Acondicionados.docx"
    ),
    Path(
        r"C:\Users\marco\OneDrive - Caja Costarricense de Seguro Social\2025"
        r"\05-Mayo-2025\Aires Acondicionados"
        r"\GF-EIM-IT-0019-2025---Esp. Tec. Aires Acondicionados.docx"
    ),
    Path(
        r"C:\Users\marco\OneDrive - Caja Costarricense de Seguro Social\2024"
        r"\09-Setiembre-2024\Aires Acondicionados"
        r"\GF-EIM-IT-0053-2024---Esp. Tec. Aires Acondicionados.docx"
    ),
    Path(
        r"C:\Users\marco\OneDrive - Caja Costarricense de Seguro Social\2023"
        r"\06-Junio-2023"
        r"\GF-EIM-IT-0020-2023-Esp Téc Mante AA Central 19-06-2023.docx"
    ),
    Path(
        r"C:\Users\marco\OneDrive - Caja Costarricense de Seguro Social\2023"
        r"\01-Enero-2023\Aires Acondicionados Norte"
        r"\GF-EIM-IT-0001-2023---Esp. Tec. AA Norte.docx"
    ),
]

OUTPUT = Path(__file__).resolve().parent / "index" / "docx_sources.json"


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def inspect(path: Path) -> dict[str, object]:
    doc = Document(str(path))
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": paragraph.text,
                "runs": [
                    {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                    }
                    for run in paragraph.runs
                    if run.text
                ],
            }
        )

    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row_index, row in enumerate(table.rows):
            rows.append(
                {
                    "row": row_index,
                    "cells": [
                        "\n".join(
                            paragraph.text
                            for paragraph in cell.paragraphs
                            if paragraph.text.strip()
                        )
                        for cell in row.cells
                    ],
                }
            )
        tables.append({"index": table_index, "rows": rows})

    return {
        "path": str(path),
        "filename": path.name,
        "exists": path.exists(),
        "sha256": file_hash(path),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    result = [inspect(path) for path in SOURCES]
    OUTPUT.write_text(
        json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(
        json.dumps(
            [
                {
                    "filename": item["filename"],
                    "paragraph_count": item["paragraph_count"],
                    "table_count": item["table_count"],
                }
                for item in result
            ],
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
