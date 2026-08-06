import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORK = Path(r"C:\Users\marco\OneDrive - Caja Costarricense de Seguro Social\Documentos\FreeCAD\Macros")
DATA_PATH = WORK / "tmp" / "ac_specs" / "technical_content.json"
OUT_DIR = Path(
    r"C:\Users\marco\OneDrive - Caja Costarricense de Seguro Social\EIMGF"
    r"\Recopilacion_Especificaciones_Aire_Acondicionado"
)
OUT_PATH = OUT_DIR / "Especificaciones_tecnicas_consolidadas_aire_acondicionado.docx"

PAGE_DXA = 12240
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "173553"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "2C7A73"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "DDEFEA"
LIGHT_GRAY = "F2F4F7"
MUTED = "5B6573"
INK = "1E293B"
WHITE = "FFFFFF"
GOLD = "7A5A00"
GOLD_FILL = "FFF4CC"
RED = "9B1C1C"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        tag = qn(f"w:{side}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C3CF", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)


def style_table_text(table, header=True, body_size=9):
    for r_idx, row in enumerate(table.rows):
        prevent_row_split(row)
        if r_idx == 0 and header:
            set_repeat_table_header(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=body_size if r_idx else body_size,
                        color=WHITE if (header and r_idx == 0) else INK,
                        bold=(header and r_idx == 0),
                    )


def set_cell_text(cell, text, *, bold=False, size=9, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), "18")
    r_pr.append(size_el)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_numbering_definition(document, fmt="bullet"):
    numbering = document.part.numbering_part.element
    existing_abs = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    existing_num = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(existing_abs, default=-1) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if fmt == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if fmt == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    if fmt == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_bullet(doc, text, num_id=None, *, bold_prefix=None, style_name="List Bullet"):
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if num_id is not None:
        apply_num(p, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, size=11, color=INK, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc, title, body, fill=LIGHT_TEAL, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=accent, size=8)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    r1 = p1.add_run(title)
    set_run_font(r1, size=10.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=INK)
    return table


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep_with_next(p)
    return p


def clean_text(text):
    return (
        str(text)
        .replace(">= 17", "≥ 17")
        .replace(">=17", "≥ 17")
        .replace("<= ", "≤ ")
        .replace("\u00a0", " ")
    )


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    patch_builtin_list_style(doc, "List Bullet")
    patch_builtin_list_style(doc, "List Number")


def patch_builtin_list_style(doc, style_name):
    style = doc.styles[style_name]
    p_pr = style._element.pPr
    if p_pr is None:
        return
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        return
    num_id = num_id_el.get(qn("w:val"))
    numbering = doc.part.numbering_part.element
    num = next(
        (x for x in numbering.findall(qn("w:num")) if x.get(qn("w:numId")) == num_id),
        None,
    )
    if num is None:
        return
    abs_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
    abstract = next(
        (
            x
            for x in numbering.findall(qn("w:abstractNum"))
            if x.get(qn("w:abstractNumId")) == abs_id
        ),
        None,
    )
    if abstract is None:
        return
    lvl = next(
        (x for x in abstract.findall(qn("w:lvl")) if x.get(qn("w:ilvl")) == "0"),
        None,
    )
    if lvl is None:
        return
    lvl_p_pr = lvl.find(qn("w:pPr"))
    if lvl_p_pr is None:
        lvl_p_pr = OxmlElement("w:pPr")
        lvl.append(lvl_p_pr)
    ind = lvl_p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        lvl_p_pr.append(ind)
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    tabs = lvl_p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        lvl_p_pr.append(tabs)
    for old in list(tabs):
        tabs.remove(old)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    spacing = lvl_p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        lvl_p_pr.append(spacing)
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CCSS · Gerencia Financiera · Infraestructura y Mantenimiento")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7488, 1872], indent_dxa=0)
    for cell in table.row_cells(0):
        set_cell_margins(cell, {"top": 0, "bottom": 0, "start": 0, "end": 0})
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    r = left.add_run("Especificaciones técnicas consolidadas de aire acondicionado")
    set_run_font(r, size=8.5, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    label = right.add_run("Página ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_field(right)


def add_title_block(doc, data):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ESPECIFICACIONES TÉCNICAS")
    set_run_font(r, size=10.5, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(clean_text(data["title"]))
    set_run_font(r, size=25, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(clean_text(data["subtitle"]))
    set_run_font(r, size=13, color=MUTED)

    rows = [
        ("Estado", data["status"]),
        ("Corte documental y normativo", data["cutoff_date"]),
        ("Unidad usuaria", "Equipo de Infraestructura y Mantenimiento · Gerencia Financiera"),
        ("Aplicación", "Compra, instalación, puesta en marcha, garantía y mantenimiento"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table, color="CAD3DD", size=6)
    for idx, (label, value) in enumerate(rows):
        shade_cell(table.cell(idx, 0), LIGHT_BLUE)
        set_cell_text(table.cell(idx, 0), label, bold=True, size=9.5, color=NAVY)
        set_cell_text(table.cell(idx, 1), clean_text(value), size=9.5)
    for row in table.rows:
        prevent_row_split(row)

    doc.add_paragraph()
    add_callout(
        doc,
        "Condición de uso",
        "Documento técnico consolidado para revisión institucional. Antes de incorporarlo a un cartel o contrato deberán completarse las cantidades, capacidades, ubicaciones, planos, condiciones del sitio y criterios de evaluación del procedimiento específico.",
        fill=GOLD_FILL,
        accent=GOLD,
    )


def add_control_documental(doc, data, bullet_num):
    add_heading(doc, "Control documental", 1)
    p = doc.add_paragraph(
        "La consolidación conserva como autoridad el PDF firmado y utiliza el archivo editable únicamente para recuperar la estructura y facilitar la actualización."
    )
    p.paragraph_format.space_after = Pt(7)
    for source in data["base_documents"]:
        text = (
            f"{source['role']}: {source['document']} "
            f"({source['date']}). {source['authority']}."
        )
        add_bullet(doc, clean_text(text), bullet_num)


def add_scope_and_use(doc, data, bullet_num):
    add_heading(doc, "Alcance de esta versión", 1)
    p = doc.add_paragraph(clean_text(data["scope_note"]))
    p.paragraph_format.space_after = Pt(7)
    points = [
        "Las cláusulas se redactan como requisitos mínimos; el oferente podrá proponer soluciones superiores cuando sean compatibles con el diseño, la normativa y la infraestructura existente.",
        "La especificación particular de cada contratación prevalece en cantidades, ubicaciones, capacidades, horarios y restricciones de ejecución.",
        "No se incluyen chillers ni unidades paquete. Los sistemas VRF se aplicarán únicamente cuando el cuadro de equipos y el diseño particular los indiquen.",
    ]
    for item in points:
        add_bullet(doc, item, bullet_num)


def add_contents(doc, sections):
    add_heading(doc, "Contenido", 1)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [900, 8460])
    set_table_borders(table, color="D5DDE6", size=5)
    set_cell_text(table.cell(0, 0), "N.º", bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), "Sección", bold=True, size=9, color=WHITE)
    shade_cell(table.cell(0, 0), NAVY)
    shade_cell(table.cell(0, 1), NAVY)
    for section in sections:
        heading = clean_text(section["heading"])
        match = re.match(r"^(\d+)\.\s*(.*)$", heading)
        num, title = (match.group(1), match.group(2)) if match else ("", heading)
        cells = table.add_row().cells
        set_cell_text(cells[0], num, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], title, size=9)
        if int(num or 0) % 2 == 0:
            shade_cell(cells[0], LIGHT_GRAY)
            shade_cell(cells[1], LIGHT_GRAY)
    style_table_text(table, header=True, body_size=9)


def add_requirements(doc, sections):
    doc.add_page_break()
    add_heading(doc, "Especificaciones consolidadas", 1)
    p = doc.add_paragraph(
        "Los siguientes requisitos se aplicarán en conjunto con el cuadro particular de equipos, los planos, la visita técnica y las instrucciones del fabricante."
    )
    p.paragraph_format.space_after = Pt(10)
    for section in sections:
        add_heading(doc, clean_text(section["heading"]), 2)
        for requirement in section["requirements"]:
            add_bullet(
                doc,
                clean_text(requirement),
                style_name="List Number",
            )


def add_schedule_template(doc):
    add_heading(doc, "Cuadro particular que debe completar cada procedimiento", 1)
    p = doc.add_paragraph(
        "Este cuadro funciona como control mínimo de definición. Puede sustituirse por planos o anexos más detallados, siempre que conserven la misma información."
    )
    p.style = doc.styles["Normal"]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table, color="B8C3CF", size=6)
    set_cell_text(table.cell(0, 0), "Dato", bold=True, size=9.5, color=WHITE)
    set_cell_text(table.cell(0, 1), "Definición requerida", bold=True, size=9.5, color=WHITE)
    shade_cell(table.cell(0, 0), NAVY)
    shade_cell(table.cell(0, 1), NAVY)
    rows = [
        ("Ubicación e identificación", "Edificio, nivel, recinto, código del equipo y ubicación de unidades interior/exterior."),
        ("Tipo y cantidad", "Configuración requerida y cantidad por línea."),
        ("Capacidad y desempeño", "Carga de diseño o capacidad nominal; rango de operación; REEE y demás indicadores exigibles."),
        ("Servicio eléctrico", "Tensión, fases, frecuencia, disponibilidad del circuito y restricciones de recorrido."),
        ("Tuberías y drenaje", "Longitudes estimadas, desniveles, punto de descarga, bomba si aplica y pases previstos."),
        ("Distribución de aire", "Difusores, retornos, ductos, balance y condiciones de ventilación cuando aplique."),
        ("Ambiente", "Corrosividad, exposición solar/lluvia, atmósfera marina, polvo, humedad o uso crítico."),
        ("Ejecución", "Horarios, permisos, acceso, izaje, trabajos en altura, protección de ocupantes y acabados."),
        ("Equipos existentes", "Retiro, recuperación de refrigerante, traslado, disposición o entrega a la Administración."),
        ("Recepción", "Pruebas particulares, responsables, plazo, capacitación y documentos exigidos."),
    ]
    for label, detail in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=9)
        set_cell_text(cells[1], detail, size=9)
    style_table_text(table, header=True, body_size=9)


def add_maintenance_routine(doc, routine):
    doc.add_page_break()
    add_heading(doc, "Rutina mínima de mantenimiento preventivo", 1)
    p = doc.add_paragraph(
        "La rutina se ejecutará por unidad interior y unidad exterior, con trazabilidad común del sistema. Cada actividad deberá registrarse con condición encontrada, acción ejecutada, medición cuando corresponda y resultado."
    )
    p.paragraph_format.space_after = Pt(8)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [720, 6840, 1800])
    set_table_borders(table, color="B8C3CF", size=6)
    headers = ["N.º", "Actividad mínima", "Registro"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else None)
        shade_cell(table.cell(0, idx), NAVY)
    for idx, item in enumerate(routine, start=1):
        cells = table.add_row().cells
        set_cell_text(cells[0], idx, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], clean_text(item), size=9)
        set_cell_text(cells[2], "Conforme / Hallazgo / N/A", size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
        if idx % 2 == 0:
            for cell in cells:
                shade_cell(cell, LIGHT_GRAY)
    style_table_text(table, header=True, body_size=9)


def add_acceptance_checklist(doc, bullet_num):
    add_heading(doc, "Entregables mínimos para recepción", 1)
    items = [
        "Acta de puesta en marcha por sistema, firmada por el responsable técnico del contratista.",
        "Registro de modelo y serie de unidades interior y exterior, controles y accesorios principales.",
        "Certificados y fichas técnicas del conjunto ofrecido, incluidos eficiencia energética y compatibilidad de componentes.",
        "Resultados de pruebas eléctricas, refrigerantes, drenaje, operación, temperaturas, flujo o balance según el tipo de sistema.",
        "Planos o croquis conforme a obra, con circuitos, tuberías, drenajes, ductos, controles y puntos de desconexión.",
        "Manual de operación y mantenimiento, garantías, calendario de servicio y contactos de atención.",
        "Inventario fotográfico del antes, durante y después, incluyendo trabajos ocultos antes de cerrar cielos, ductos o acabados.",
        "Constancia de capacitación al personal designado por la Administración.",
        "Comprobantes de recuperación y gestión de refrigerante y residuos cuando se retiren equipos existentes.",
        "Lista de pendientes cerrada o, excepcionalmente, plan de corrección aprobado sin afectar seguridad ni operación esencial.",
    ]
    for item in items:
        add_bullet(doc, item, bullet_num)


def add_normative_references(doc, sources):
    doc.add_page_break()
    add_heading(doc, "Normativa y referencias técnicas", 1)
    p = doc.add_paragraph(
        "Se aplicará la versión vigente a la fecha de publicación del procedimiento. Cuando una norma citada sea sustituida, se utilizará su edición vigente o la disposición que legalmente la reemplace, salvo justificación técnica expresa."
    )
    p.paragraph_format.space_after = Pt(8)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2500, 5360, 1500])
    set_table_borders(table, color="B8C3CF", size=6)
    for idx, header in enumerate(["Instrumento / enlace", "Aplicación en estas especificaciones", "Condición"]):
        set_cell_text(table.cell(0, idx), header, bold=True, size=8.5, color=WHITE)
        shade_cell(table.cell(0, idx), NAVY)
    for idx, source in enumerate(sources, start=1):
        cells = table.add_row().cells
        cells[0].text = ""
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r = p0.add_run(clean_text(source["instrument"]))
        set_run_font(r, size=8.5, color=INK, bold=True)
        p_link = cells[0].add_paragraph()
        p_link.paragraph_format.space_after = Pt(0)
        add_hyperlink(p_link, "Fuente oficial", source["url"])
        set_cell_text(cells[1], clean_text(source["application"]), size=8.5)
        set_cell_text(cells[2], clean_text(source["status"]), size=8.5)
        if idx % 2 == 0:
            for cell in cells:
                shade_cell(cell, LIGHT_GRAY)
    style_table_text(table, header=True, body_size=8.5)


def add_closing_note(doc):
    doc.add_paragraph()
    add_callout(
        doc,
        "Revisión previa a publicación",
        "La unidad técnica deberá verificar el cuadro de equipos, la carga térmica o selección, el servicio eléctrico disponible, la ruta de drenaje, las condiciones de corrosión, el acceso para mantenimiento, la documentación de eficiencia energética y la compatibilidad del refrigerante antes de aprobar cada procedimiento.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )


def audit_document(doc):
    errors = []
    section = doc.sections[0]
    values = {
        "page_width": section.page_width.twips,
        "page_height": section.page_height.twips,
        "top_margin": section.top_margin.twips,
        "bottom_margin": section.bottom_margin.twips,
        "left_margin": section.left_margin.twips,
        "right_margin": section.right_margin.twips,
        "header_distance": section.header_distance.twips,
        "footer_distance": section.footer_distance.twips,
    }
    expected = {
        "page_width": 12240,
        "page_height": 15840,
        "top_margin": 1440,
        "bottom_margin": 1440,
        "left_margin": 1440,
        "right_margin": 1440,
        "header_distance": 708,
        "footer_distance": 708,
    }
    for key, exp in expected.items():
        if abs(values[key] - exp) > 2:
            errors.append(f"{key}={values[key]}, expected {exp}")

    for idx, table in enumerate(doc.tables, start=1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        grid_sum = sum(
            int(col.get(qn("w:w"))) for col in table._tbl.tblGrid.findall(qn("w:gridCol"))
        )
        if tbl_w is None or int(tbl_w.get(qn("w:w"))) != CONTENT_DXA:
            errors.append(f"table {idx}: tblW")
        expected_indent = 0 if table._tbl.getparent() is section.footer._element else TABLE_INDENT_DXA
        if tbl_ind is None:
            errors.append(f"table {idx}: tblInd missing")
        if grid_sum != CONTENT_DXA:
            errors.append(f"table {idx}: grid sum {grid_sum}")
        for r_idx, row in enumerate(table.rows):
            widths = []
            for cell in row.cells:
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                widths.append(int(tc_w.get(qn("w:w"))) if tc_w is not None else -1)
            if sum(widths) != CONTENT_DXA:
                errors.append(f"table {idx} row {r_idx}: tcW sum {sum(widths)}")
                break
    if errors:
        raise RuntimeError("DOCX token audit failed:\n" + "\n".join(errors[:25]))


def build():
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        configure_header_footer(section)

    add_title_block(doc, data)
    add_control_documental(doc, data, None)
    add_scope_and_use(doc, data, None)
    add_contents(doc, data["sections"])
    add_requirements(doc, data["sections"])
    add_schedule_template(doc)
    add_maintenance_routine(doc, data["maintenance_routine"])
    add_acceptance_checklist(doc, None)
    add_normative_references(doc, data["normative_sources"])
    add_closing_note(doc)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    audit_document(doc)
    doc.save(OUT_PATH)
    print(
        json.dumps(
            {
                "output": str(OUT_PATH),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(data["sections"]),
                "requirements": sum(len(s["requirements"]) for s in data["sections"]),
                "maintenance_items": len(data["maintenance_routine"]),
                "normative_sources": len(data["normative_sources"]),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    build()
